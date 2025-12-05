import os
import re
from docx import Document
from docx.shared import Pt

def convert_markdown_to_docx(source_dir):
    for filename in os.listdir(source_dir):
        if filename.endswith(".md"):
            md_path = os.path.join(source_dir, filename)
            docx_path = os.path.join(source_dir, filename.replace(".md", ".docx"))
            
            print(f"Converting {filename} to DOCX...")
            
            doc = Document()
            
            with open(md_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
                
            in_code_block = False
            
            for line in lines:
                line = line.rstrip()
                
                # Handle Code Blocks
                if line.startswith("```"):
                    in_code_block = not in_code_block
                    continue
                
                if in_code_block:
                    p = doc.add_paragraph(line)
                    p.style = 'No Spacing'
                    font = p.runs[0].font
                    font.name = 'Courier New'
                    font.size = Pt(10)
                    continue
                
                # Handle Headers
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("#### "):
                    doc.add_heading(line[5:], level=4)
                
                # Handle Lists
                elif line.strip().startswith("- ") or line.strip().startswith("* "):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif re.match(r"^\d+\. ", line.strip()):
                    doc.add_paragraph(line.strip().split(". ", 1)[1], style='List Number')
                    
                # Handle Normal Text (skip empty lines if desired, or keep them)
                elif line.strip() == "":
                    # doc.add_paragraph("") # Optional: Add empty paragraph for spacing
                    pass
                else:
                    doc.add_paragraph(line)
            
            doc.save(docx_path)
            print(f"Saved {docx_path}")

if __name__ == "__main__":
    guides_dir = os.path.join(os.getcwd(), "docs", "guides")
    if os.path.exists(guides_dir):
        convert_markdown_to_docx(guides_dir)
    else:
        print(f"Directory not found: {guides_dir}")
